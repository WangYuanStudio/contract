<<<<<<< HEAD
# -*- coding: utf-8 -*-

import sys, os, datetime
from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from aip import AipOcr
from docx import Document


class App(QWidget):
    def __init__(self):
        """
        类初始化
        """
        super().__init__()

        # 设置窗体属性
        self.setWindowTitle(u'廖老师合同自动归档工具')
        self.setObjectName('App')
        self.setWindowIcon(QIcon('favicon.ico'))
        self.screen = QDesktopWidget().screenGeometry()                     # 获取屏幕尺寸
        self.screen.setHeight(1020)
        self.setGeometry(0, 0, self.screen.width(), self.screen.height())   # 使窗口大小于屏幕大小一致
        self.setWindowFlags(Qt.WindowMinimizeButtonHint |                   # 使能最小化按钮
                            Qt.WindowCloseButtonHint)                       # 使能关闭按钮
        self.setFixedSize(self.width(), self.height())                      # 固定窗体大小

        # 初始化类字段
        self.units = []
        self.projects = {'count': 0, 'items': [], 'index': {'global': 0, 'local': 0, 'unit': 0, 'sure': 0}}
        self.stop_work = False

        self.img_file = ''
        self.img_file_new = ''
        self.img_rect = QRect(328, 8, self.screen.width()-352, self.screen.height()-216)
        self.img_area_rect = QRect(0, 0, 0, 0)
        self.img_displacement = 0

        # 设置全局样式
        self.setStyleSheet(
            'QWidget#App{background-color:#666}'
            'QPushButton{font-size:14px;color:#aaa;border:1px solid #888; background-color:#555}'
            'QPushButton:hover{color:#bbb;border:1px solid #999}'
            'QPushButton:pressed{margin:1px}'
            'QPushButton:disabled{background-color:#444;color:#666;border:1px solid #666}'
            'QFrame{background-color:#444;border-right:2px solid #aaa;border-top:2px solid #aaa}'
            'QLineEdit{background-color:#444;color:#aaa;border:1px solid #888;padding:0 8px}'
            'QLineEdit:focus{color:#eee;border:1px solid #ccc}'
            'QLineEdit:disabled{color:#666;border:1px solid #666}'
            'QFrame QLabel{color:#aaa;font-size:12px;border:none}'
            'QMessageBox QLabel{color:black;background-color:white;font-size:18px;padding:32px 40px;border:none}'
            'QMessageBox QPushButton{font-size:14px;color:black;background-color:#eee;border:1px solid #69d;padding:8px 16px}'
            'QMessageBox QPushButton:hover{color:black;border:1px solid #48c}'
            'QMessageBox QPushButton:focus{border:1px solid #04c}'
            'QScrollBar:vertical {background-color: #555;width: 8px;}'
            'QScrollBar::handle:vertical {background-color:#888;min-height:24px;}'
            'QScrollBar::handle:vertical:hover {background-color:#aaa;min-height:24px;}'
            'QScrollBar::add-line:vertical {height:0;}'
            'QScrollBar::sub-line:vertical {height:0;}'
            'QScrollBar::add-page:vertical,QScrollBar::sub-page:vertical{background-color:rgba(0,0,0,0%)}'
        )

        # 显示组件
        self.show_cmd()
        self.show_controller()
        self.show_setting()
        self.show_editor()
        self.show_scrollbar()

        # 导入保存的记录
        self.ins()

    def ins(self):
        if os.path.exists('log/history.json'):
            with open('log/history.json', 'r') as fp:
                txt = fp.read()
                self.projects = eval(txt)

                if self.projects['count']:
                    self.btn_ocr.setEnabled(True)
                    self.btn_edit.setEnabled(True)
                    self.btn_save.setEnabled(True)
                    self.btn_generate.setEnabled(True)

                    self.cmd.append(self.helper_color('读取到历史记录！', 'successful'))

    def show_cmd(self):
        """
        显示“命令行”
        :return:
        """
        self.cmd = QTextEdit(self)
        self.cmd.setObjectName('CMD')
        self.cmd.setGeometry(320, self.screen.height()-200, self.screen.width()-320, 200)
        self.cmd.setReadOnly(True)
        self.cmd.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.cmd.textChanged.connect(self.event_cmd_text_changed)
        self.cmd.setStyleSheet('QTextEdit#CMD{font-size:12px; background-color:#444; color:#aaa; '
                               'border:none; border-top:2px solid #aaa;padding:8px}')
        self.cmd.setText(
            '┌───────────────────────────┐\n'
            '│ 合同自动归档程序 - 终端\n'
            '├───────────────────────────┤\n'
            '│ 版本：1.0\n'
            '│ 日期：2018-07-03\n'
            '└───────────────────────────┘\n\n'
        )
        self.cmd.show()

    def show_controller(self):
        """
        显示“控制栏”
        :return:
        """
        self.frame_con = QFrame(self)
        self.frame_con.setObjectName('FrameCon')
        self.frame_con.setGeometry(0, 0, 320, 200)
        self.frame_con.setStyleSheet('QFrame#FrameCon{border-top:none}')

        self.btn_scan = QPushButton('扫描图片', self.frame_con)
        self.btn_scan.setGeometry(32, 32, 124, 40)
        self.btn_scan.clicked.connect(self.event_btn_scan_clicked)

        self.btn_ocr = QPushButton('识别合同', self.frame_con)
        self.btn_ocr.setGeometry(164, 32, 124, 40)
        self.btn_ocr.setEnabled(False)
        self.btn_ocr.clicked.connect(self.event_btn_ocr_clicked)

        self.btn_edit = QPushButton('编辑信息', self.frame_con)
        self.btn_edit.setGeometry(32, 80, 124, 40)
        self.btn_edit.setEnabled(False)
        self.btn_edit.clicked.connect(self.event_btn_edit_clicked)

        self.btn_generate = QPushButton('创建归档', self.frame_con)
        self.btn_generate.setGeometry(164, 80, 124, 40)
        self.btn_generate.setEnabled(False)
        self.btn_generate.clicked.connect(self.event_btn_generate_clicked)

        self.btn_save = QPushButton('保存记录', self.frame_con)
        self.btn_save.setGeometry(32, 128, 124, 40)
        self.btn_save.clicked.connect(self.event_btn_save_clicked)

        self.btn_exit = QPushButton('退出', self.frame_con)
        self.btn_exit.setGeometry(164, 128, 58, 40)
        self.btn_exit.clicked.connect(self.close)

        self.btn_min = QPushButton('最小化', self.frame_con)
        self.btn_min.setGeometry(230, 128, 58, 40)
        self.btn_min.clicked.connect(self.showMinimized)

    def show_setting(self):
        """
        显示“配置栏”
        :return:
        """
        self.frame_set = QFrame(self)
        self.frame_set.setGeometry(0, 200, 320, 154)

        self.lab1 = QLabel('合同图片目录：', self.frame_set)
        self.lab1.setGeometry(32, 32, 120, 24)

        self.edit_origin = QLineEdit('origins', self.frame_set)
        self.edit_origin.setGeometry(120, 32, 168, 24)

        self.lab2 = QLabel('归档存放目录：', self.frame_set)
        self.lab2.setGeometry(32, 64, 120, 24)

        self.edit_result = QLineEdit('results', self.frame_set)
        self.edit_result.setGeometry(120, 64, 168, 24)

        self.lab3 = QLabel('记录存放目录：', self.frame_set)
        self.lab3.setGeometry(32, 96, 120, 24)

        self.edit_log = QLineEdit('log', self.frame_set)
        self.edit_log.setGeometry(120, 96, 168, 24)

    def show_editor(self):
        """
        显示“编辑栏”
        :return:
        """
        self.frame_edi = QFrame(self)
        self.frame_edi.setGeometry(0, 354, 320, self.screen.height()-354)

        self.lab_position = QLabel('位置：' + self.helper_color('全局[0/0]    当前[0/0]    已确认[0]', 'disabled'), self.frame_edi)
        self.lab_position.setGeometry(32, 32, 256, 24)

        self.lab_unit = QLabel('单位：' + self.helper_color('Unknown', 'disabled'), self.frame_edi)
        self.lab_unit.setGeometry(32, 56, 256, 24)

        self.lab_status = QLabel('状态：' + self.helper_color('Unknown', 'disabled'), self.frame_edi)
        self.lab_status.setGeometry(32, 80, 256, 24)

        self.lab_origin = QLabel('图片：' + self.helper_color('Unknown', 'disabled'), self.frame_edi)
        self.lab_origin.setGeometry(32, 104, 256, 24)

        self.lab5 = QLabel('合同日期：', self.frame_edi)
        self.lab5.setGeometry(32, 144, 256, 24)

        self.edit_date = QLineEdit(self.frame_edi)
        self.edit_date.setGeometry(32, 168, 256, 24)
        self.edit_date.setPlaceholderText('Enter 键完成输入')
        self.edit_date.setEnabled(False)

        self.lab6 = QLabel('合同名称：', self.frame_edi)
        self.lab6.setGeometry(32, 200, 256, 24)

        self.edit_name = QTextEdit(self.frame_edi)
        self.edit_name.setObjectName('EditName')
        self.edit_name.setGeometry(32, 232, 256, 80)
        self.edit_name.setPlaceholderText('Ctrl+Enter 键确认信息')
        self.edit_name.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.edit_name.setEnabled(False)
        self.edit_name.setStyleSheet(
            'QTextEdit#EditName{border:1px solid #888;color:#aaa}'
            'QTextEdit#EditName:focus{border:1px solid #ccc;color:#eee}'
            'QTextEdit#EditName:disabled{color:#666;border:1px solid #666}'
        )

        self.btn_last = QPushButton('上一个', self.frame_edi)
        self.btn_last.setGeometry(32, self.frame_edi.height()-64, 80, 32)
        self.btn_last.clicked.connect(self.event_btn_last_clicked)
        self.btn_last.setEnabled(False)

        self.btn_next = QPushButton('下一个', self.frame_edi)
        self.btn_next.setGeometry(120, self.frame_edi.height()-64, 80, 32)
        self.btn_next.clicked.connect(self.event_btn_next_clicked)
        self.btn_next.setEnabled(False)

        self.btn_sure = QPushButton('确认', self.frame_edi)
        self.btn_sure.setGeometry(208, self.frame_edi.height()-64, 80, 32)
        self.btn_sure.clicked.connect(self.event_btn_sure_clicked)
        self.btn_sure.setEnabled(False)

    def show_scrollbar(self):
        """
        创建照片显示区的滚动条
        :return:
        """
        self.scrollbar = QScrollBar(self)
        self.scrollbar.setGeometry(self.screen.width()-16, 8, 8, self.screen.height()-216)
        self.scrollbar.valueChanged.connect(self.event_scrollbar_value_changed)
        self.scrollbar.hide()

    def event_btn_scan_clicked(self):
        """
        “扫描合同图片”按钮被单击事件
        :return:
        """
        if self.btn_scan.text() == '扫描图片':
            self.stop_work = False
            self.handler_disable_all_btn()
            self.btn_scan.setEnabled(True)
            self.btn_scan.setText('取消扫描')
            self.handler_scan()
            QApplication.processEvents()
            self.btn_scan.setText('扫描图片')
        else:
            self.btn_scan.setText('扫描图片')
            self.stop_work = True
            self.units = []
            self.projects = {'count': 0}
            self.cmd.append(self.helper_color('<br>扫描被取消！', 'warning'))

    def event_btn_ocr_clicked(self):
        """
        “识别合同信息”按钮被单击事件
        :return:
        """
        if self.btn_ocr.text() == '识别合同':
            self.stop_work = False
            self.handler_disable_all_btn()
            self.btn_ocr.setEnabled(True)
            self.btn_generate.setEnabled(False)

            self.btn_ocr.setText('取消识别')
            self.handler_ocr()
            QApplication.processEvents()
            self.btn_ocr.setText('识别合同')
        else:
            self.btn_ocr.setText('识别合同')
            self.stop_work = True
            self.handler_disable_all_btn()
            self.btn_scan.setEnabled(True)
            self.btn_ocr.setEnabled(True)
            self.cmd.append(self.helper_color('<br>识别被取消！', 'warning'))
            if self.projects['count'] != 0:
                for project in self.projects['items']:
                    data = project['data']
                    for item in data:
                        item['ocr_text'] = ''
                        item['result'] = {}

    def event_btn_save_clicked(self):
        """
        “保存程序记录”按钮被单击
        :return:
        """
        self.handler_save()

    def event_btn_edit_clicked(self):
        """
        "编辑识别结果"按钮被单击事件
        :return:
        """
        if self.btn_edit.text() == '编辑信息':
            self.handler_disable_all_btn()
            self.btn_edit.setEnabled(True)
            self.btn_edit.setText('完成编辑')
            self.cmd.append('\n正在编辑合同信息...')
            self.handler_show_info()
        else:
            self.handler_disable_all_btn(True)
            self.btn_edit.setText('编辑信息')
            self.handler_disable_editor()
            self.cmd.append(self.helper_color('已完成合同信息的编辑！', 'successful'))

    def event_btn_generate_clicked(self):
        """
        “创建合同归档”按钮被单击事件
        :return:
        """
        if self.btn_generate.text() == '创建归档':
            self.stop_work = False
            self.handler_disable_all_btn()
            self.btn_generate.setEnabled(True)
            self.btn_generate.setText('取消创建')
            self.handler_generate()
            QApplication.processEvents()
            self.handler_disable_all_btn(True)
            self.btn_generate.setText('创建归档')
        else:
            self.stop_work = True
            self.btn_generate.setText('创建归档')
            self.cmd.append(self.helper_color('<br>创建归档被取消！', 'warning'))

    def event_btn_clean_clicked(self):
        """
        “清空所有归档”按钮被单击
        :return:
        """
        self.handler_clean()

    def event_btn_last_clicked(self):
        """
        “上一个”按钮被单击事件
        :return:
        """
        if self.projects['index']['global'] != 0:
            self.projects['index']['global'] -= 1

            if self.projects['index']['local'] != 0:
                self.projects['index']['local'] -= 1
            else:
                self.projects['index']['unit'] -= 1
                self.projects['index']['local'] = self.projects['items'][self.projects['index']['unit']]['count'] - 1

        self.handler_show_info()

    def event_btn_next_clicked(self):
        """
        “下一个”按钮被单击事件
        :return:
        """
        if self.projects['index']['global'] != self.projects['count'] - 1:
            self.projects['index']['global'] += 1

            if self.projects['index']['local'] != self.projects['items'][self.projects['index']['unit']]['count'] - 1:
                self.projects['index']['local'] += 1
            else:
                self.projects['index']['local'] = 0
                self.projects['index']['unit'] += 1

        self.handler_show_info()

    def event_btn_sure_clicked(self):
        """
        “确认”按钮被单击事件
        :return:
        """
        status = self.projects['items'][self.projects['index']['unit']]['data'][self.projects['index']['local']]['modified']
        if not status:
            sure_count = self.projects['index']['sure']
            self.projects['index']['sure'] = sure_count + 1

        self.projects['items'][self.projects['index']['unit']]['data'][self.projects['index']['local']]['modified'] = True
        self.projects['items'][self.projects['index']['unit']]['data'][self.projects['index']['local']]['result']['name'] = self.edit_name.toPlainText()
        self.projects['items'][self.projects['index']['unit']]['data'][self.projects['index']['local']]['result']['date'] = self.edit_date.text()

        self.event_btn_next_clicked()

    def event_cmd_text_changed(self):
        """
        “命令行”文本框内容被改变事件
        :return:
        """
        self.cmd.verticalScrollBar().setSliderPosition(self.cmd.verticalScrollBar().maximum())

    def event_scrollbar_value_changed(self):
        """
        照片显示区滚动条值被改变事件
        :return:
        """
        disp = 0 - self.scrollbar.value() * 20
        min = 0 - (self.img_area_rect.height() - self.img_rect.height())
        if self.scrollbar.value() == self.scrollbar.maximum():
            disp = min
        self.img_displacement = disp
        self.update()

    def closeEvent(self, event):
        """
        重写 closeEvent 事件
        :param event:
        :return: None
        """
        reply = QMessageBox.warning(self, '提示', "确定退出程序？", QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if reply == QMessageBox.Yes:
            event.accept()
        else:
            event.ignore()

    def paintEvent(self, event):
        """
        重写 paintEvent 事件
        :param event:
        :return:
        """
        # 如果图片文件有更新，则绘制新图片
        if self.img_file != self.img_file_new:
            self.img_file = self.img_file_new

            image = QPixmap()
            image.load(self.img_file)
            painter = QPainter(self)

            width = image.width()
            height = image.height()

            if width > self.img_rect.width():
                scale = width/height
                width = self.img_rect.width()
                height = int(width/scale)

            lef = int(self.img_rect.width()/2 - width/2)

            if height > self.img_rect.height():
                self.scrollbar.setMaximum(int((height - self.img_rect.height())/20))
                self.scrollbar.setValue(0)
                self.scrollbar.show()
            else:
                self.scrollbar.hide()

            self.img_area_rect.setRect(self.img_rect.left() + lef, self.img_rect.top() + 0, width, height)
            self.img_displacement = 0
            self.img_rotate = 0

            painter.drawPixmap(self.img_area_rect, image)

        # 否则重绘移位或旋转后的图像
        else:
            image = QPixmap()
            image.load(self.img_file)
            painter = QPainter(self)

            width = image.width()
            height = image.height()

            if width > self.img_rect.width():
                scale = width / height
                width = self.img_rect.width()
                height = int(width / scale)

            lef = int(self.img_rect.width()/2 - width/2) + self.img_rect.left()
            top = self.img_rect.top() + self.img_displacement

            self.img_area_rect.setRect(lef, top, width, height)
            painter.drawPixmap(self.img_area_rect, image)

    def wheelEvent(self, event):
        """
        重写 wheelEvent 事件
        :param event:
        :return:
        """
        if self.img_area_rect.height() <= self.img_rect.height():
            return

        x = event.x()
        y = event.y()
        l = self.img_rect.left()
        r = self.img_rect.right()
        t = self.img_rect.top()
        b = self.img_rect.bottom()
        if x >= l and x <= r and y >= t and y <= b:
            angle = event.angleDelta()
            value = angle.y()
            disp = self.img_displacement + value
            min = 0 - (self.img_area_rect.height() - self.img_rect.height())
            if disp < min:
                disp = min
            elif disp > 0:
                disp = 0

            self.scrollbar.setValue((0-disp)/20)
            self.img_displacement = disp
            self.update()

    def keyPressEvent(self, event):
        """
        重写 keyPressEvent 事件
        :param event:
        :return:
        """
        event = QKeyEvent(event)
        if self.edit_name.hasFocus() and event.key() == 16777220:
            self.event_btn_sure_clicked()
        elif self.edit_date.hasFocus() and event.key() == 16777220:  # 16777220 回车键
            self.edit_name.setFocus()
            self.edit_name.selectAll()

    def handler_disable_all_btn(self, enabled=False):
        """
        处理程序：禁用所有主要控制按钮
        :return:
        """
        self.btn_scan.setEnabled(enabled)
        self.btn_ocr.setEnabled(enabled)
        self.btn_edit.setEnabled(enabled)
        self.btn_generate.setEnabled(enabled)
        self.btn_scan.setText('扫描图片')
        self.btn_ocr.setText('识别合同')
        self.btn_edit.setText('编辑信息')
        self.btn_generate.setText('创建归档')

    def handler_disable_editor(self, enabled=False):
        """
        处理程序：禁用所有编辑区控件
        :param enabled:
        :return:
        """
        self.btn_last.setEnabled(enabled)
        self.btn_next.setEnabled(enabled)
        self.btn_sure.setEnabled(enabled)
        self.edit_name.setEnabled(enabled)
        self.edit_date.setEnabled(enabled)

    def handler_scan(self):
        """
        处理程序：扫描合同单位及对应单位下的所有合同图片
        :return:
        """
        # 清空记录
        self.units = []
        self.projects = {'count': 0, 'items': [], 'index': {'global': 0, 'local': 0, 'unit': 0, 'sure': 0}}

        # 获取单位分类
        self.cmd.append('\n正在扫描合同单位...')
        # root = os.path.abspath(self.edit_origin.text())
        root = self.edit_origin.text()
        if root[-1] != '/' and root[-1] != '\\':
            root += '/'

        if not os.path.exists(root):
            self.cmd.append(self.helper_color('合同图片目录不存在，扫描终止！', 'danger'))
            return

        path = os.listdir(root)
        tmp_text = "已扫描到单位："
        for p in path:
            # 中途取消扫描
            if self.stop_work:
                return

            if os.path.isdir(root + p):
                self.units.append(p)
                tmp_text += self.helper_color(p, 'importance') + '  '

        self.cmd.append(tmp_text)

        # 扫描各个单位下的合同图片
        self.cmd.append('\n正在扫描各个单位下的合同图片...')
        extension = ['bmp', 'png', 'jpg']
        projects = []
        for t in self.units:
            # 中途取消扫描
            if self.stop_work:
                return

            root_tmp = root + t + '/'
            path = os.listdir(root_tmp)
            path.sort()
            data = []
            for p in path:
                # 中途取消扫描
                if self.stop_work:
                    return

                ext = p[-3:]
                if ext in extension:
                    data.append({'ocr_text': '', 'result': {'name': '', 'date': ''}, 'origin': root_tmp + p, 'modified': False})
                    self.cmd.append(self.helper_color(os.path.abspath(root_tmp + p), 'disabled'))

            if len(data) != 0:
                project = {'unit': t, 'count': len(data), 'doc': '', 'data': data}
                projects.append(project)
                self.cmd.append('共扫描到 ' + self.helper_color(t, 'info') +
                                ' 的合同图片 ' + self.helper_color(str(len(data)), 'info') + ' 张\n\n\n\n\n')

        if len(projects) == 0:
            self.cmd.append(self.helper_color('<br>扫描程序已完成，但未扫描到任何合同图片！', 'warning'))
        else:
            count = 0
            for project in projects:
                count += project['count']
            self.projects['count'] = count
            self.projects['items'] = projects
            self.cmd.append(self.helper_color('<br>扫描程序已完成！', 'successful'))
            self.btn_ocr.setEnabled(True)
            self.btn_edit.setEnabled(True)
            self.btn_generate.setEnabled(True)

    def handler_ocr(self):
        """
        处理程序：OCR 文本识别
        :return:
        """
        if self.projects['count'] == 0:
            self.cmd.append(self.helper_color('<br>无合同图片，不能进行识别！', 'warning'))
            return

        # 连接百度云
        self.cmd.append('\n正在连接到百度云 OCR 接口...')
        QApplication.processEvents()

        app_id = "11407546"
        api_key = "lY6vQLFc1zMBCotBThZWgEPO"
        secret_key = "FhWvLVonrxbtKa4bKFgVLAFnsMyRKbTU"
        template_id1 = "2a9be4f38c0806b9f5dfd3c3616b7560"
        # template_id2 = "f61002a4a009e77a0cd53775b64a771b"

        client = AipOcr(app_id, api_key, secret_key)
        client.setConnectionTimeoutInMillis(20000)
        client.setSocketTimeoutInMillis(30000)

        # 文本识别
        for project in self.projects['items']:
            # 中途取消扫描
            if self.stop_work:
                return

            index = 0
            self.cmd.append('<br>正在识别 ' + self.helper_color(project['unit'], 'info') + ' 的合同图片...')
            QApplication.processEvents()

            for data in project['data']:
                # 中途取消扫描
                if self.stop_work:
                    return

                index += 1
                data['result'] = {'name': '', 'date': ''}
                data['modified'] = False
                try:
                    with open(data['origin'], 'rb') as fp:
                        image = fp.read()
                        data['ocr_text'] = client.custom(image, template_id1)
                        QApplication.processEvents()

                        if data['ocr_text']['error_code'] != 0:
                            self.cmd.append(str(index) + '/' + str(project['count']) + self.helper_color('  识别失败: ', 'warning')
                                            + self.helper_color(os.path.abspath(data['origin']), 'disabled'))
                            error_code = data['ocr_text']['error_code']
                            error_msg = data['ocr_text']['error_msg']
                            self.cmd.append('---- 错误代码：' + str(error_code))
                            self.cmd.append('---- 错误描述：' + str(error_msg))
                        else:
                            data['result'] = {
                                'name': data['ocr_text']['data']['ret'][1]['word'],
                                'date': data['ocr_text']['data']['ret'][0]['word']
                            }

                            self.cmd.append(str(index) + '/' + str(project['count']) + self.helper_color('  识别成功：', 'successful')
                                + self.helper_color(os.path.abspath(data['origin']), 'disabled'))
                            self.cmd.append('---- 合同名称：' + self.helper_color(data['result']['name'], 'importance'))
                            self.cmd.append('---- 合同日期：' + data['result']['date'])
                except:
                    self.cmd.append(str(index) + '/' + str(project['count']) + self.helper_color('  识别出错: ', 'danger')
                        + os.path.abspath(data['origin']))

                QApplication.processEvents()

        self.cmd.append(self.helper_color('<br>OCR 识别程序已完成！', 'successful'))
        self.handler_disable_all_btn(True)

    def handler_show_info(self):
        """
        处理程序：显示合同信息
        :return:
        """
        # print(self.projects)
        index_unit = self.projects['index']['unit']
        index_global = self.projects['index']['global']
        index_local = self.projects['index']['local']
        count_sure = self.projects['index']['sure']
        count_global = self.projects['count']
        count_local = self.projects['items'][index_unit]['count']
        unit = self.projects['items'][index_unit]['unit']
        status = self.projects['items'][index_unit]['data'][index_local]['modified']
        origin = self.projects['items'][index_unit]['data'][index_local]['origin']
        name = self.projects['items'][index_unit]['data'][index_local]['result']['name']
        date = self.projects['items'][index_unit]['data'][index_local]['result']['date']

        self.lab_position.setText('位置：全局[{0}/{1}]  当前[{2}/{3}]  已确认[{4}]'.format(index_global+1, count_global, index_local+1, count_local, count_sure))
        self.lab_unit.setText('单位：' + self.helper_color(unit, 'info'))
        self.lab_origin.setText('图片：' + str.split(origin, '/')[-1])
        self.edit_name.setText(name)
        self.edit_date.setText(date)
        if status:
            self.lab_status.setText('状态：' + self.helper_color('已确认', 'successful'))
        else:
            self.lab_status.setText('状态：' + self.helper_color('未确认', 'warning'))

        self.handler_disable_editor(True)
        if index_global == 0:
            self.btn_last.setEnabled(False)
        if index_global == count_global:
            self.btn_next.setEnabled(False)

        self.edit_date.setFocus()
        self.edit_date.selectAll()

        self.img_file_new = origin
        self.cmd.append(self.helper_color(origin, 'disabled'))
        self.update()

    def handler_save(self):
        """
        处理程序：保存 OCR 识别结果和 cmd 输出日志
        :return:
        """
        root = os.path.abspath(self.edit_log.text())
        if root[-1] != '/' and root[-1] != '\\':
            root += '/'
        if not os.path.exists(root):
            self.cmd.append(self.helper_color('<br>输出日志目录不存在，正在创建目录:', 'warning') + self.helper_color(root, 'disabled'))
            try:
                os.mkdir(root)
            except:
                self.cmd.append(self.helper_color('创建目录失败，识别程序被终止！', 'danger') )
                return

        # filename1 = root + datetime.datetime.now().strftime('compact_console_log_%Y%m%d%H%M%S.html')
        # filename2 = root + datetime.datetime.now().strftime('compact_ocr_%Y%m%d%H%M%S.json')
        filename1 = root + 'history.html'
        filename2 = root + 'history.json'

        self.cmd.append('\n正在保存控制台输出日志...')
        with open(filename1, 'w') as fp:
            html = self.cmd.toHtml()
            html += '<style type="text/css">body{background-color:#444} p{color:#aaa}</style>'
            res = fp.write(html)
        if res:
            self.cmd.append(self.helper_color('成功保存文件：', 'successful') + self.helper_color(filename1, 'importance'))
        else:
            self.cmd.append(self.helper_color('保存文件失败！', 'danger') + self.helper_color(filename1, 'disabled'))

        self.cmd.append('\n正在保存 OCR 文本识别结果...')
        projects = self.projects.copy()
        # for project in projects['items']:
        #     data = project['data']
        #     for item in data:
        #         try:
        #             item.pop('image')
        #         except:
        #             pass
        with open(filename2, 'w') as fp:
            res = fp.write(str(projects))
        if res:
            self.cmd.append(self.helper_color('成功保存文件：', 'successful') + self.helper_color(filename2, 'importance'))
        else:
            self.cmd.append(self.helper_color('保存文件失败！', 'danger') + self.helper_color(filename2, 'disabled'))

    def handler_generate(self):
        """
        处理程序：创建合同归档
        :return:
        """
        # 目录是否存在
        root = os.path.abspath(self.edit_result.text())
        if root[-1] != '/' and root[-1] != '\\':
            root += '/'
        if not os.path.exists(root):
            self.cmd.append(self.helper_color('<br>归档存放目录不存在，正在创建目录:', 'warning') + self.helper_color(root, 'disabled'))
            try:
                os.mkdir(root)
            except:
                self.cmd.append(self.helper_color('创建目录失败，创建归档程序被终止！', 'danger'))
                return

        # 是否无内容
        if self.projects['count'] == 0:
            self.cmd.append(self.helper_color('无合同，请先扫描和识别合同！', 'warning'))
            return

        # 创建 Word 归档
        index = 0
        self.cmd.append('\n正在创建合同归档文件...')
        for project in self.projects['items']:
            # 中途取消扫描
            if self.stop_work:
                return

            QApplication.processEvents()

            index += 1

            temp_file = 'template.docx'
            if not os.path.exists(os.getcwd() + '/' + temp_file):
                self.cmd.append(self.helper_color('找不到模板 docx 文件，创建该单位的归档文件失败！', 'danger'))
                continue

            try:
                # 读取模板文档
                doc = Document(temp_file)
                tab = doc.tables[1]
                rows = tab.rows
                rows_count = len(rows)

                # 如果模板表格的行数不够
                while rows_count <= project['count']:
                    tab.add_row()
                    rows_count = len(rows)
            except:
                self.cmd.append(self.helper_color('无效的模板文件，创建该单位的归档文件失败！', 'danger'))
                continue

            QApplication.processEvents()

            # 填充数据
            i = 0
            for data in project['data']:
                # 中途取消扫描
                if self.stop_work:
                    return

                QApplication.processEvents()

                i += 1
                result = data['result']
                if result is {}:
                    self.cmd.append(self.helper_color('该合同的信息无效：', 'warning') + self.helper_color(data['origin'], 'disabled'))
                    continue

                content = [str(i), project['unit'], result['date'], result['name'], '', '永久', '']
                try:
                    row = rows[i]
                    cells = row.cells
                    for j in range(7):
                        cells[j].text = content[j]
                except:
                    self.cmd.append(self.helper_color('填充数据失败：', 'warning') + self.helper_color(result['name'], 'disabled'))
                    continue

                QApplication.processEvents()

            # 保存新文档
            filename = '合同卷内目录' + str(index) + '-' + project['unit'] + '.docx'
            try:
                doc.save(root + filename)
            except:
                self.cmd.append(self.helper_color('创建归档文件失败：', 'danger') + filename)
                continue

            QApplication.processEvents()
            if os.path.exists(root + filename):
                self.cmd.append(self.helper_color('成功创建归档文件：', 'successful') + self.helper_color(filename, 'importance'))
                project['doc'] = filename
            else:
                self.cmd.append(self.helper_color('创建归档文件失败：', 'danger') + filename)

            QApplication.processEvents()

        self.cmd.append(self.helper_color('<br>创建归档程序已完成！', 'successful'))

    @staticmethod
    def helper_color(text, type='default'):
        """
        辅助函数：获取命令行字体颜色
        :param text:
        :param type:
        :return:
        """
        if type is 'default':
            return '<span style="color:#aaa">' + text + '</span>'
        if type is 'danger':
            return '<span style="color:red">' + text + '</span>'
        if type is 'successful':
            return '<span style="color:green">' + text + '</span>'
        if type is 'warning':
            return '<span style="color:orange">' + text + '</span>'
        if type is 'active':
            return '<span style="color:#6F6">' + text + '</span>'
        if type is 'info':
            return '<span style="color:#08F">' + text + '</span>'
        if type is 'disabled':
            return '<span style="color:#555">' + text + '</span>'
        if type is 'importance':
            return '<span style="color:white">' + text + '</span>'

        return text

    @staticmethod
    def helper_get_img_bin(img_path):
        """
        辅助函数：获取图片数据
        :param img_path:
        :return:
        """
        if os.path.exists(img_path):
            with open(img_path, 'rb') as fp:
                return fp.read()
        else:
            return None

    @staticmethod
    def helper_del_files(root):
        """
        辅助函数：删除指定目录下的所有文件
        :param root:
        :return:
        """
        root = os.path.abspath(root)
        files = os.listdir(root)
        try:
            for file in files:
                tmp_path = os.path.join(root, file)
                if os.path.isdir(tmp_path):
                    App.helper_del_files(tmp_path)
                    os.removedirs(tmp_path)
                else:
                    os.remove(tmp_path)
        except:
            return False

        return os.listdir(root) == []


if __name__ == '__main__':
    app = QApplication(sys.argv)
    win = App()
    win.show()
    sys.exit(app.exec_())
=======
# -*- coding: utf-8 -*-

import sys, os, datetime
from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from aip import AipOcr
from docx import Document


class App(QWidget):
    def __init__(self):
        """
        类初始化
        """
        super().__init__()

        # 设置窗体属性
        self.setWindowTitle(u'廖老师合同自动归档工具')
        self.setObjectName('App')
        self.setWindowIcon(QIcon('favicon.ico'))
        self.screen = QDesktopWidget().screenGeometry()                     # 获取屏幕尺寸
        self.screen.setHeight(1020)
        self.setGeometry(0, 0, self.screen.width(), self.screen.height())   # 使窗口大小于屏幕大小一致
        self.setWindowFlags(Qt.WindowMinimizeButtonHint |                   # 使能最小化按钮
                            Qt.WindowCloseButtonHint)                       # 使能关闭按钮
        self.setFixedSize(self.width(), self.height())                      # 固定窗体大小

        # 初始化类字段
        self.units = []
        self.projects = {'count': 0, 'items': [], 'index': {'global': 0, 'local': 0, 'unit': 0, 'sure': 0}}
        self.stop_work = False

        self.img_file = ''
        self.img_file_new = ''
        self.img_rect = QRect(328, 8, self.screen.width()-352, self.screen.height()-216)
        self.img_area_rect = QRect(0, 0, 0, 0)
        self.img_displacement = 0

        # 设置全局样式
        self.setStyleSheet(
            'QWidget#App{background-color:#666}'
            'QPushButton{font-size:14px;color:#aaa;border:1px solid #888; background-color:#555}'
            'QPushButton:hover{color:#bbb;border:1px solid #999}'
            'QPushButton:pressed{margin:1px}'
            'QPushButton:disabled{background-color:#444;color:#666;border:1px solid #666}'
            'QFrame{background-color:#444;border-right:2px solid #aaa;border-top:2px solid #aaa}'
            'QLineEdit{background-color:#444;color:#aaa;border:1px solid #888;padding:0 8px}'
            'QLineEdit:focus{color:#eee;border:1px solid #ccc}'
            'QLineEdit:disabled{color:#666;border:1px solid #666}'
            'QFrame QLabel{color:#aaa;font-size:12px;border:none}'
            'QMessageBox QLabel{color:black;background-color:white;font-size:18px;padding:32px 40px;border:none}'
            'QMessageBox QPushButton{font-size:14px;color:black;background-color:#eee;border:1px solid #69d;padding:8px 16px}'
            'QMessageBox QPushButton:hover{color:black;border:1px solid #48c}'
            'QMessageBox QPushButton:focus{border:1px solid #04c}'
            'QScrollBar:vertical {background-color: #555;width: 8px;}'
            'QScrollBar::handle:vertical {background-color:#888;min-height:24px;}'
            'QScrollBar::handle:vertical:hover {background-color:#aaa;min-height:24px;}'
            'QScrollBar::add-line:vertical {height:0;}'
            'QScrollBar::sub-line:vertical {height:0;}'
            'QScrollBar::add-page:vertical,QScrollBar::sub-page:vertical{background-color:rgba(0,0,0,0%)}'
        )

        # 显示组件
        self.show_cmd()
        self.show_controller()
        self.show_setting()
        self.show_editor()
        self.show_scrollbar()

        # 导入保存的记录
        self.ins()

    def ins(self):
        if os.path.exists('log/history.json'):
            with open('log/history.json', 'r') as fp:
                txt = fp.read()
                self.projects = eval(txt)

                if self.projects['count']:
                    self.btn_ocr.setEnabled(True)
                    self.btn_edit.setEnabled(True)
                    self.btn_save.setEnabled(True)
                    self.btn_generate.setEnabled(True)

                    self.cmd.append(self.helper_color('读取到历史记录！', 'successful'))

    def show_cmd(self):
        """
        显示“命令行”
        :return:
        """
        self.cmd = QTextEdit(self)
        self.cmd.setObjectName('CMD')
        self.cmd.setGeometry(320, self.screen.height()-200, self.screen.width()-320, 200)
        self.cmd.setReadOnly(True)
        self.cmd.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.cmd.textChanged.connect(self.event_cmd_text_changed)
        self.cmd.setStyleSheet('QTextEdit#CMD{font-size:12px; background-color:#444; color:#aaa; '
                               'border:none; border-top:2px solid #aaa;padding:8px}')
        self.cmd.setText(
            '┌───────────────────────────┐\n'
            '│ 合同自动归档程序 - 终端\n'
            '├───────────────────────────┤\n'
            '│ 版本：1.0\n'
            '│ 日期：2018-07-03\n'
            '└───────────────────────────┘\n\n'
        )
        self.cmd.show()

    def show_controller(self):
        """
        显示“控制栏”
        :return:
        """
        self.frame_con = QFrame(self)
        self.frame_con.setObjectName('FrameCon')
        self.frame_con.setGeometry(0, 0, 320, 200)
        self.frame_con.setStyleSheet('QFrame#FrameCon{border-top:none}')

        self.btn_scan = QPushButton('扫描图片', self.frame_con)
        self.btn_scan.setGeometry(32, 32, 124, 40)
        self.btn_scan.clicked.connect(self.event_btn_scan_clicked)

        self.btn_ocr = QPushButton('识别合同', self.frame_con)
        self.btn_ocr.setGeometry(164, 32, 124, 40)
        self.btn_ocr.setEnabled(False)
        self.btn_ocr.clicked.connect(self.event_btn_ocr_clicked)

        self.btn_edit = QPushButton('编辑信息', self.frame_con)
        self.btn_edit.setGeometry(32, 80, 124, 40)
        self.btn_edit.setEnabled(False)
        self.btn_edit.clicked.connect(self.event_btn_edit_clicked)

        self.btn_generate = QPushButton('创建归档', self.frame_con)
        self.btn_generate.setGeometry(164, 80, 124, 40)
        self.btn_generate.setEnabled(False)
        self.btn_generate.clicked.connect(self.event_btn_generate_clicked)

        self.btn_save = QPushButton('保存记录', self.frame_con)
        self.btn_save.setGeometry(32, 128, 124, 40)
        self.btn_save.clicked.connect(self.event_btn_save_clicked)

        self.btn_exit = QPushButton('退出', self.frame_con)
        self.btn_exit.setGeometry(164, 128, 58, 40)
        self.btn_exit.clicked.connect(self.close)

        self.btn_min = QPushButton('最小化', self.frame_con)
        self.btn_min.setGeometry(230, 128, 58, 40)
        self.btn_min.clicked.connect(self.showMinimized)

    def show_setting(self):
        """
        显示“配置栏”
        :return:
        """
        self.frame_set = QFrame(self)
        self.frame_set.setGeometry(0, 200, 320, 154)

        self.lab1 = QLabel('合同图片目录：', self.frame_set)
        self.lab1.setGeometry(32, 32, 120, 24)

        self.edit_origin = QLineEdit('origins', self.frame_set)
        self.edit_origin.setGeometry(120, 32, 168, 24)

        self.lab2 = QLabel('归档存放目录：', self.frame_set)
        self.lab2.setGeometry(32, 64, 120, 24)

        self.edit_result = QLineEdit('results', self.frame_set)
        self.edit_result.setGeometry(120, 64, 168, 24)

        self.lab3 = QLabel('记录存放目录：', self.frame_set)
        self.lab3.setGeometry(32, 96, 120, 24)

        self.edit_log = QLineEdit('log', self.frame_set)
        self.edit_log.setGeometry(120, 96, 168, 24)

    def show_editor(self):
        """
        显示“编辑栏”
        :return:
        """
        self.frame_edi = QFrame(self)
        self.frame_edi.setGeometry(0, 354, 320, self.screen.height()-354)

        self.lab_position = QLabel('位置：' + self.helper_color('全局[0/0]    当前[0/0]    已确认[0]', 'disabled'), self.frame_edi)
        self.lab_position.setGeometry(32, 32, 256, 24)

        self.lab_unit = QLabel('单位：' + self.helper_color('Unknown', 'disabled'), self.frame_edi)
        self.lab_unit.setGeometry(32, 56, 256, 24)

        self.lab_status = QLabel('状态：' + self.helper_color('Unknown', 'disabled'), self.frame_edi)
        self.lab_status.setGeometry(32, 80, 256, 24)

        self.lab_origin = QLabel('图片：' + self.helper_color('Unknown', 'disabled'), self.frame_edi)
        self.lab_origin.setGeometry(32, 104, 256, 24)

        self.lab5 = QLabel('合同日期：', self.frame_edi)
        self.lab5.setGeometry(32, 144, 256, 24)

        self.edit_date = QLineEdit(self.frame_edi)
        self.edit_date.setGeometry(32, 168, 256, 24)
        self.edit_date.setPlaceholderText('Enter 键完成输入')
        self.edit_date.setEnabled(False)

        self.lab6 = QLabel('合同名称：', self.frame_edi)
        self.lab6.setGeometry(32, 200, 256, 24)

        self.edit_name = QTextEdit(self.frame_edi)
        self.edit_name.setObjectName('EditName')
        self.edit_name.setGeometry(32, 232, 256, 80)
        self.edit_name.setPlaceholderText('Ctrl+Enter 键确认信息')
        self.edit_name.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.edit_name.setEnabled(False)
        self.edit_name.setStyleSheet(
            'QTextEdit#EditName{border:1px solid #888;color:#aaa}'
            'QTextEdit#EditName:focus{border:1px solid #ccc;color:#eee}'
            'QTextEdit#EditName:disabled{color:#666;border:1px solid #666}'
        )

        self.btn_last = QPushButton('上一个', self.frame_edi)
        self.btn_last.setGeometry(32, self.frame_edi.height()-64, 80, 32)
        self.btn_last.clicked.connect(self.event_btn_last_clicked)
        self.btn_last.setEnabled(False)

        self.btn_next = QPushButton('下一个', self.frame_edi)
        self.btn_next.setGeometry(120, self.frame_edi.height()-64, 80, 32)
        self.btn_next.clicked.connect(self.event_btn_next_clicked)
        self.btn_next.setEnabled(False)

        self.btn_sure = QPushButton('确认', self.frame_edi)
        self.btn_sure.setGeometry(208, self.frame_edi.height()-64, 80, 32)
        self.btn_sure.clicked.connect(self.event_btn_sure_clicked)
        self.btn_sure.setEnabled(False)

    def show_scrollbar(self):
        """
        创建照片显示区的滚动条
        :return:
        """
        self.scrollbar = QScrollBar(self)
        self.scrollbar.setGeometry(self.screen.width()-16, 8, 8, self.screen.height()-216)
        self.scrollbar.valueChanged.connect(self.event_scrollbar_value_changed)
        self.scrollbar.hide()

    def event_btn_scan_clicked(self):
        """
        “扫描合同图片”按钮被单击事件
        :return:
        """
        if self.btn_scan.text() == '扫描图片':
            self.stop_work = False
            self.handler_disable_all_btn()
            self.btn_scan.setEnabled(True)
            self.btn_scan.setText('取消扫描')
            self.handler_scan()
            QApplication.processEvents()
            self.btn_scan.setText('扫描图片')
        else:
            self.btn_scan.setText('扫描图片')
            self.stop_work = True
            self.units = []
            self.projects = {'count': 0}
            self.cmd.append(self.helper_color('<br>扫描被取消！', 'warning'))

    def event_btn_ocr_clicked(self):
        """
        “识别合同信息”按钮被单击事件
        :return:
        """
        if self.btn_ocr.text() == '识别合同':
            self.stop_work = False
            self.handler_disable_all_btn()
            self.btn_ocr.setEnabled(True)
            self.btn_generate.setEnabled(False)

            self.btn_ocr.setText('取消识别')
            self.handler_ocr()
            QApplication.processEvents()
            self.btn_ocr.setText('识别合同')
        else:
            self.btn_ocr.setText('识别合同')
            self.stop_work = True
            self.handler_disable_all_btn()
            self.btn_scan.setEnabled(True)
            self.btn_ocr.setEnabled(True)
            self.cmd.append(self.helper_color('<br>识别被取消！', 'warning'))
            if self.projects['count'] != 0:
                for project in self.projects['items']:
                    data = project['data']
                    for item in data:
                        item['ocr_text'] = ''
                        item['result'] = {}

    def event_btn_save_clicked(self):
        """
        “保存程序记录”按钮被单击
        :return:
        """
        self.handler_save()

    def event_btn_edit_clicked(self):
        """
        "编辑识别结果"按钮被单击事件
        :return:
        """
        if self.btn_edit.text() == '编辑信息':
            self.handler_disable_all_btn()
            self.btn_edit.setEnabled(True)
            self.btn_edit.setText('完成编辑')
            self.cmd.append('\n正在编辑合同信息...')
            self.handler_show_info()
        else:
            self.handler_disable_all_btn(True)
            self.btn_edit.setText('编辑信息')
            self.handler_disable_editor()
            self.cmd.append(self.helper_color('已完成合同信息的编辑！', 'successful'))

    def event_btn_generate_clicked(self):
        """
        “创建合同归档”按钮被单击事件
        :return:
        """
        if self.btn_generate.text() == '创建归档':
            self.stop_work = False
            self.handler_disable_all_btn()
            self.btn_generate.setEnabled(True)
            self.btn_generate.setText('取消创建')
            self.handler_generate()
            QApplication.processEvents()
            self.handler_disable_all_btn(True)
            self.btn_generate.setText('创建归档')
        else:
            self.stop_work = True
            self.btn_generate.setText('创建归档')
            self.cmd.append(self.helper_color('<br>创建归档被取消！', 'warning'))

    def event_btn_clean_clicked(self):
        """
        “清空所有归档”按钮被单击
        :return:
        """
        self.handler_clean()

    def event_btn_last_clicked(self):
        """
        “上一个”按钮被单击事件
        :return:
        """
        if self.projects['index']['global'] != 0:
            self.projects['index']['global'] -= 1

            if self.projects['index']['local'] != 0:
                self.projects['index']['local'] -= 1
            else:
                self.projects['index']['unit'] -= 1
                self.projects['index']['local'] = self.projects['items'][self.projects['index']['unit']]['count'] - 1

        self.handler_show_info()

    def event_btn_next_clicked(self):
        """
        “下一个”按钮被单击事件
        :return:
        """
        if self.projects['index']['global'] != self.projects['count'] - 1:
            self.projects['index']['global'] += 1

            if self.projects['index']['local'] != self.projects['items'][self.projects['index']['unit']]['count'] - 1:
                self.projects['index']['local'] += 1
            else:
                self.projects['index']['local'] = 0
                self.projects['index']['unit'] += 1

        self.handler_show_info()

    def event_btn_sure_clicked(self):
        """
        “确认”按钮被单击事件
        :return:
        """
        status = self.projects['items'][self.projects['index']['unit']]['data'][self.projects['index']['local']]['modified']
        if not status:
            sure_count = self.projects['index']['sure']
            self.projects['index']['sure'] = sure_count + 1

        self.projects['items'][self.projects['index']['unit']]['data'][self.projects['index']['local']]['modified'] = True
        self.projects['items'][self.projects['index']['unit']]['data'][self.projects['index']['local']]['result']['name'] = self.edit_name.toPlainText()
        self.projects['items'][self.projects['index']['unit']]['data'][self.projects['index']['local']]['result']['date'] = self.edit_date.text()

        self.event_btn_next_clicked()

    def event_cmd_text_changed(self):
        """
        “命令行”文本框内容被改变事件
        :return:
        """
        self.cmd.verticalScrollBar().setSliderPosition(self.cmd.verticalScrollBar().maximum())

    def event_scrollbar_value_changed(self):
        """
        照片显示区滚动条值被改变事件
        :return:
        """
        disp = 0 - self.scrollbar.value() * 20
        min = 0 - (self.img_area_rect.height() - self.img_rect.height())
        if self.scrollbar.value() == self.scrollbar.maximum():
            disp = min
        self.img_displacement = disp
        self.update()

    def closeEvent(self, event):
        """
        重写 closeEvent 事件
        :param event:
        :return: None
        """
        reply = QMessageBox.warning(self, '提示', "确定退出程序？", QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if reply == QMessageBox.Yes:
            event.accept()
        else:
            event.ignore()

    def paintEvent(self, event):
        """
        重写 paintEvent 事件
        :param event:
        :return:
        """
        # 如果图片文件有更新，则绘制新图片
        if self.img_file != self.img_file_new:
            self.img_file = self.img_file_new

            image = QPixmap()
            image.load(self.img_file)
            painter = QPainter(self)

            width = image.width()
            height = image.height()

            if width > self.img_rect.width():
                scale = width/height
                width = self.img_rect.width()
                height = int(width/scale)

            lef = int(self.img_rect.width()/2 - width/2)

            if height > self.img_rect.height():
                self.scrollbar.setMaximum(int((height - self.img_rect.height())/20))
                self.scrollbar.setValue(0)
                self.scrollbar.show()
            else:
                self.scrollbar.hide()

            self.img_area_rect.setRect(self.img_rect.left() + lef, self.img_rect.top() + 0, width, height)
            self.img_displacement = 0
            self.img_rotate = 0

            painter.drawPixmap(self.img_area_rect, image)

        # 否则重绘移位或旋转后的图像
        else:
            image = QPixmap()
            image.load(self.img_file)
            painter = QPainter(self)

            width = image.width()
            height = image.height()

            if width > self.img_rect.width():
                scale = width / height
                width = self.img_rect.width()
                height = int(width / scale)

            lef = int(self.img_rect.width()/2 - width/2) + self.img_rect.left()
            top = self.img_rect.top() + self.img_displacement

            self.img_area_rect.setRect(lef, top, width, height)
            painter.drawPixmap(self.img_area_rect, image)

    def wheelEvent(self, event):
        """
        重写 wheelEvent 事件
        :param event:
        :return:
        """
        if self.img_area_rect.height() <= self.img_rect.height():
            return

        x = event.x()
        y = event.y()
        l = self.img_rect.left()
        r = self.img_rect.right()
        t = self.img_rect.top()
        b = self.img_rect.bottom()
        if x >= l and x <= r and y >= t and y <= b:
            angle = event.angleDelta()
            value = angle.y()
            disp = self.img_displacement + value
            min = 0 - (self.img_area_rect.height() - self.img_rect.height())
            if disp < min:
                disp = min
            elif disp > 0:
                disp = 0

            self.scrollbar.setValue((0-disp)/20)
            self.img_displacement = disp
            self.update()

    def keyPressEvent(self, event):
        """
        重写 keyPressEvent 事件
        :param event:
        :return:
        """
        event = QKeyEvent(event)
        if self.edit_name.hasFocus() and event.key() == 16777220:
            self.event_btn_sure_clicked()
        elif self.edit_date.hasFocus() and event.key() == 16777220:  # 16777220 回车键
            self.edit_name.setFocus()
            self.edit_name.selectAll()

    def handler_disable_all_btn(self, enabled=False):
        """
        处理程序：禁用所有主要控制按钮
        :return:
        """
        self.btn_scan.setEnabled(enabled)
        self.btn_ocr.setEnabled(enabled)
        self.btn_edit.setEnabled(enabled)
        self.btn_generate.setEnabled(enabled)
        self.btn_scan.setText('扫描图片')
        self.btn_ocr.setText('识别合同')
        self.btn_edit.setText('编辑信息')
        self.btn_generate.setText('创建归档')

    def handler_disable_editor(self, enabled=False):
        """
        处理程序：禁用所有编辑区控件
        :param enabled:
        :return:
        """
        self.btn_last.setEnabled(enabled)
        self.btn_next.setEnabled(enabled)
        self.btn_sure.setEnabled(enabled)
        self.edit_name.setEnabled(enabled)
        self.edit_date.setEnabled(enabled)

    def handler_scan(self):
        """
        处理程序：扫描合同单位及对应单位下的所有合同图片
        :return:
        """
        # 清空记录
        self.units = []
        self.projects = {'count': 0, 'items': [], 'index': {'global': 0, 'local': 0, 'unit': 0, 'sure': 0}}

        # 获取单位分类
        self.cmd.append('\n正在扫描合同单位...')
        # root = os.path.abspath(self.edit_origin.text())
        root = self.edit_origin.text()
        if root[-1] != '/' and root[-1] != '\\':
            root += '/'

        if not os.path.exists(root):
            self.cmd.append(self.helper_color('合同图片目录不存在，扫描终止！', 'danger'))
            return

        path = os.listdir(root)
        tmp_text = "已扫描到单位："
        for p in path:
            # 中途取消扫描
            if self.stop_work:
                return

            if os.path.isdir(root + p):
                self.units.append(p)
                tmp_text += self.helper_color(p, 'importance') + '  '

        self.cmd.append(tmp_text)

        # 扫描各个单位下的合同图片
        self.cmd.append('\n正在扫描各个单位下的合同图片...')
        extension = ['bmp', 'png', 'jpg']
        projects = []
        for t in self.units:
            # 中途取消扫描
            if self.stop_work:
                return

            root_tmp = root + t + '/'
            path = os.listdir(root_tmp)
            path.sort()
            data = []
            for p in path:
                # 中途取消扫描
                if self.stop_work:
                    return

                ext = p[-3:]
                if ext in extension:
                    data.append({'ocr_text': '', 'result': {'name': '', 'date': ''}, 'origin': root_tmp + p, 'modified': False})
                    self.cmd.append(self.helper_color(os.path.abspath(root_tmp + p), 'disabled'))

            if len(data) != 0:
                project = {'unit': t, 'count': len(data), 'doc': '', 'data': data}
                projects.append(project)
                self.cmd.append('共扫描到 ' + self.helper_color(t, 'info') +
                                ' 的合同图片 ' + self.helper_color(str(len(data)), 'info') + ' 张\n\n\n\n\n')

        if len(projects) == 0:
            self.cmd.append(self.helper_color('<br>扫描程序已完成，但未扫描到任何合同图片！', 'warning'))
        else:
            count = 0
            for project in projects:
                count += project['count']
            self.projects['count'] = count
            self.projects['items'] = projects
            self.cmd.append(self.helper_color('<br>扫描程序已完成！', 'successful'))
            self.btn_ocr.setEnabled(True)
            self.btn_edit.setEnabled(True)
            self.btn_generate.setEnabled(True)

    def handler_ocr(self):
        """
        处理程序：OCR 文本识别
        :return:
        """
        if self.projects['count'] == 0:
            self.cmd.append(self.helper_color('<br>无合同图片，不能进行识别！', 'warning'))
            return

        # 连接百度云
        self.cmd.append('\n正在连接到百度云 OCR 接口...')
        QApplication.processEvents()

        app_id = "11407546"
        api_key = "lY6vQLFc1zMBCotBThZWgEPO"
        secret_key = "FhWvLVonrxbtKa4bKFgVLAFnsMyRKbTU"
        template_id1 = "2a9be4f38c0806b9f5dfd3c3616b7560"
        # template_id2 = "f61002a4a009e77a0cd53775b64a771b"

        client = AipOcr(app_id, api_key, secret_key)
        client.setConnectionTimeoutInMillis(20000)
        client.setSocketTimeoutInMillis(30000)

        # 文本识别
        for project in self.projects['items']:
            # 中途取消扫描
            if self.stop_work:
                return

            index = 0
            self.cmd.append('<br>正在识别 ' + self.helper_color(project['unit'], 'info') + ' 的合同图片...')
            QApplication.processEvents()

            for data in project['data']:
                # 中途取消扫描
                if self.stop_work:
                    return

                index += 1
                data['result'] = {'name': '', 'date': ''}
                data['modified'] = False
                try:
                    with open(data['origin'], 'rb') as fp:
                        image = fp.read()
                        data['ocr_text'] = client.custom(image, template_id1)
                        QApplication.processEvents()

                        if data['ocr_text']['error_code'] != 0:
                            self.cmd.append(str(index) + '/' + str(project['count']) + self.helper_color('  识别失败: ', 'warning')
                                            + self.helper_color(os.path.abspath(data['origin']), 'disabled'))
                            error_code = data['ocr_text']['error_code']
                            error_msg = data['ocr_text']['error_msg']
                            self.cmd.append('---- 错误代码：' + str(error_code))
                            self.cmd.append('---- 错误描述：' + str(error_msg))
                        else:
                            data['result'] = {
                                'name': data['ocr_text']['data']['ret'][1]['word'],
                                'date': data['ocr_text']['data']['ret'][0]['word']
                            }

                            self.cmd.append(str(index) + '/' + str(project['count']) + self.helper_color('  识别成功：', 'successful')
                                + self.helper_color(os.path.abspath(data['origin']), 'disabled'))
                            self.cmd.append('---- 合同名称：' + self.helper_color(data['result']['name'], 'importance'))
                            self.cmd.append('---- 合同日期：' + data['result']['date'])
                except:
                    self.cmd.append(str(index) + '/' + str(project['count']) + self.helper_color('  识别出错: ', 'danger')
                        + os.path.abspath(data['origin']))

                QApplication.processEvents()

        self.cmd.append(self.helper_color('<br>OCR 识别程序已完成！', 'successful'))
        self.handler_disable_all_btn(True)

    def handler_show_info(self):
        """
        处理程序：显示合同信息
        :return:
        """
        # print(self.projects)
        index_unit = self.projects['index']['unit']
        index_global = self.projects['index']['global']
        index_local = self.projects['index']['local']
        count_sure = self.projects['index']['sure']
        count_global = self.projects['count']
        count_local = self.projects['items'][index_unit]['count']
        unit = self.projects['items'][index_unit]['unit']
        status = self.projects['items'][index_unit]['data'][index_local]['modified']
        origin = self.projects['items'][index_unit]['data'][index_local]['origin']
        name = self.projects['items'][index_unit]['data'][index_local]['result']['name']
        date = self.projects['items'][index_unit]['data'][index_local]['result']['date']

        self.lab_position.setText('位置：全局[{0}/{1}]  当前[{2}/{3}]  已确认[{4}]'.format(index_global+1, count_global, index_local+1, count_local, count_sure))
        self.lab_unit.setText('单位：' + self.helper_color(unit, 'info'))
        self.lab_origin.setText('图片：' + str.split(origin, '/')[-1])
        self.edit_name.setText(name)
        self.edit_date.setText(date)
        if status:
            self.lab_status.setText('状态：' + self.helper_color('已确认', 'successful'))
        else:
            self.lab_status.setText('状态：' + self.helper_color('未确认', 'warning'))

        self.handler_disable_editor(True)
        if index_global == 0:
            self.btn_last.setEnabled(False)
        if index_global == count_global:
            self.btn_next.setEnabled(False)

        self.edit_date.setFocus()
        self.edit_date.selectAll()

        self.img_file_new = origin
        self.cmd.append(self.helper_color(origin, 'disabled'))
        self.update()

    def handler_save(self):
        """
        处理程序：保存 OCR 识别结果和 cmd 输出日志
        :return:
        """
        root = os.path.abspath(self.edit_log.text())
        if root[-1] != '/' and root[-1] != '\\':
            root += '/'
        if not os.path.exists(root):
            self.cmd.append(self.helper_color('<br>输出日志目录不存在，正在创建目录:', 'warning') + self.helper_color(root, 'disabled'))
            try:
                os.mkdir(root)
            except:
                self.cmd.append(self.helper_color('创建目录失败，识别程序被终止！', 'danger') )
                return

        # filename1 = root + datetime.datetime.now().strftime('compact_console_log_%Y%m%d%H%M%S.html')
        # filename2 = root + datetime.datetime.now().strftime('compact_ocr_%Y%m%d%H%M%S.json')
        filename1 = root + 'history.html'
        filename2 = root + 'history.json'

        self.cmd.append('\n正在保存控制台输出日志...')
        with open(filename1, 'w') as fp:
            html = self.cmd.toHtml()
            html += '<style type="text/css">body{background-color:#444} p{color:#aaa}</style>'
            res = fp.write(html)
        if res:
            self.cmd.append(self.helper_color('成功保存文件：', 'successful') + self.helper_color(filename1, 'importance'))
        else:
            self.cmd.append(self.helper_color('保存文件失败！', 'danger') + self.helper_color(filename1, 'disabled'))

        self.cmd.append('\n正在保存 OCR 文本识别结果...')
        projects = self.projects.copy()
        # for project in projects['items']:
        #     data = project['data']
        #     for item in data:
        #         try:
        #             item.pop('image')
        #         except:
        #             pass
        with open(filename2, 'w') as fp:
            res = fp.write(str(projects))
        if res:
            self.cmd.append(self.helper_color('成功保存文件：', 'successful') + self.helper_color(filename2, 'importance'))
        else:
            self.cmd.append(self.helper_color('保存文件失败！', 'danger') + self.helper_color(filename2, 'disabled'))

    def handler_generate(self):
        """
        处理程序：创建合同归档
        :return:
        """
        # 目录是否存在
        root = os.path.abspath(self.edit_result.text())
        if root[-1] != '/' and root[-1] != '\\':
            root += '/'
        if not os.path.exists(root):
            self.cmd.append(self.helper_color('<br>归档存放目录不存在，正在创建目录:', 'warning') + self.helper_color(root, 'disabled'))
            try:
                os.mkdir(root)
            except:
                self.cmd.append(self.helper_color('创建目录失败，创建归档程序被终止！', 'danger'))
                return

        # 是否无内容
        if self.projects['count'] == 0:
            self.cmd.append(self.helper_color('无合同，请先扫描和识别合同！', 'warning'))
            return

        # 创建 Word 归档
        index = 0
        self.cmd.append('\n正在创建合同归档文件...')
        for project in self.projects['items']:
            # 中途取消扫描
            if self.stop_work:
                return

            QApplication.processEvents()

            index += 1

            temp_file = 'template.docx'
            if not os.path.exists(os.getcwd() + '/' + temp_file):
                self.cmd.append(self.helper_color('找不到模板 docx 文件，创建该单位的归档文件失败！', 'danger'))
                continue

            try:
                # 读取模板文档
                doc = Document(temp_file)
                tab = doc.tables[1]
                rows = tab.rows
                rows_count = len(rows)

                # 如果模板表格的行数不够
                while rows_count <= project['count']:
                    tab.add_row()
                    rows_count = len(rows)
            except:
                self.cmd.append(self.helper_color('无效的模板文件，创建该单位的归档文件失败！', 'danger'))
                continue

            QApplication.processEvents()

            # 填充数据
            i = 0
            for data in project['data']:
                # 中途取消扫描
                if self.stop_work:
                    return

                QApplication.processEvents()

                i += 1
                result = data['result']
                if result is {}:
                    self.cmd.append(self.helper_color('该合同的信息无效：', 'warning') + self.helper_color(data['origin'], 'disabled'))
                    continue

                content = [str(i), project['unit'], result['date'], result['name'], '', '永久', '']
                try:
                    row = rows[i]
                    cells = row.cells
                    for j in range(7):
                        cells[j].text = content[j]
                except:
                    self.cmd.append(self.helper_color('填充数据失败：', 'warning') + self.helper_color(result['name'], 'disabled'))
                    continue

                QApplication.processEvents()

            # 保存新文档
            filename = '合同卷内目录' + str(index) + '-' + project['unit'] + '.docx'
            try:
                doc.save(root + filename)
            except:
                self.cmd.append(self.helper_color('创建归档文件失败：', 'danger') + filename)
                continue

            QApplication.processEvents()
            if os.path.exists(root + filename):
                self.cmd.append(self.helper_color('成功创建归档文件：', 'successful') + self.helper_color(filename, 'importance'))
                project['doc'] = filename
            else:
                self.cmd.append(self.helper_color('创建归档文件失败：', 'danger') + filename)

            QApplication.processEvents()

        self.cmd.append(self.helper_color('<br>创建归档程序已完成！', 'successful'))

    @staticmethod
    def helper_color(text, type='default'):
        """
        辅助函数：获取命令行字体颜色
        :param text:
        :param type:
        :return:
        """
        if type is 'default':
            return '<span style="color:#aaa">' + text + '</span>'
        if type is 'danger':
            return '<span style="color:red">' + text + '</span>'
        if type is 'successful':
            return '<span style="color:green">' + text + '</span>'
        if type is 'warning':
            return '<span style="color:orange">' + text + '</span>'
        if type is 'active':
            return '<span style="color:#6F6">' + text + '</span>'
        if type is 'info':
            return '<span style="color:#08F">' + text + '</span>'
        if type is 'disabled':
            return '<span style="color:#555">' + text + '</span>'
        if type is 'importance':
            return '<span style="color:white">' + text + '</span>'

        return text

    @staticmethod
    def helper_get_img_bin(img_path):
        """
        辅助函数：获取图片数据
        :param img_path:
        :return:
        """
        if os.path.exists(img_path):
            with open(img_path, 'rb') as fp:
                return fp.read()
        else:
            return None

    @staticmethod
    def helper_del_files(root):
        """
        辅助函数：删除指定目录下的所有文件
        :param root:
        :return:
        """
        root = os.path.abspath(root)
        files = os.listdir(root)
        try:
            for file in files:
                tmp_path = os.path.join(root, file)
                if os.path.isdir(tmp_path):
                    App.helper_del_files(tmp_path)
                    os.removedirs(tmp_path)
                else:
                    os.remove(tmp_path)
        except:
            return False

        return os.listdir(root) == []


if __name__ == '__main__':
    app = QApplication(sys.argv)
    win = App()
    win.show()
    sys.exit(app.exec_())
>>>>>>> 40105774c2483077b535794ac6d14e2c3c2c0088
